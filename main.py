"""
Аудио-Скрайбер — ИИ-ассистент для расшифровки, анализа и протоколирования аудио/видео.
"""

import requests, json, re, logging, subprocess, time, asyncio
from pathlib import Path
from io import BytesIO
from datetime import datetime
from collections import Counter

from config import *

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes

import ffmpeg
from faster_whisper import WhisperModel
from openai import OpenAI

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import networkx as nx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMP_DIR = Path("./temp")
TEMP_DIR.mkdir(exist_ok=True)

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

logger.info("Загружаю Whisper (base)...")
whisper_model = WhisperModel("base", device="cpu", compute_type="int8")
logger.info("Whisper готов.")

medium_model = None

llm_client = OpenAI(
    api_key=YANDEX_API_KEY,
    base_url=YANDEX_BASE_URL,
    default_headers={"x-folder-id": YANDEX_FOLDER_ID}
)

text_storage = {}
stats_storage = {}
result_storage = {}
audio_storage = {}
video_storage = {}
pending_cut = {}
segments_storage = {}
protocol_storage = {}
usage_stats = {"total_processed": 0, "total_seconds": 0}

def main_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("✨ Новый анализ", callback_data="send_media")],
        [InlineKeyboardButton("ℹ️ Как это работает", callback_data="about"),
         InlineKeyboardButton("🆘 Помощь", callback_data="help")],
    ])

def back_button():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("◀️ Назад в меню", callback_data="main_menu")]
    ])

def cut_or_full_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🎬 Анализировать полностью", callback_data="analyze_full")],
        [InlineKeyboardButton("✂️ Обрезать фрагмент", callback_data="cut_video")],
        [InlineKeyboardButton("◀️ Назад", callback_data="main_menu")],
    ])

def after_result_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📄 Скачать протокол (.docx)", callback_data="download_protocol")],
        [InlineKeyboardButton("⏱ Тайм-коды тем", callback_data="show_timestamps"),
         InlineKeyboardButton("📊 Статистика речи", callback_data="show_stats")],
        [InlineKeyboardButton("🎵 Аудиодорожка", callback_data="extract_audio_only"),
         InlineKeyboardButton("🔄 Показать отчёт", callback_data="show_result")],
        [InlineKeyboardButton("📝 Скачать текст (.txt)", callback_data="download_txt")],
        [InlineKeyboardButton("🌍 Перевести текст", callback_data="translate_menu")],
        [InlineKeyboardButton("✨ Новый анализ", callback_data="send_media")],
    ])

def translate_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🇷🇺 Русский", callback_data="translate_ru"),
         InlineKeyboardButton("🇬🇧 English", callback_data="translate_en"),
         InlineKeyboardButton("🇧🇾 Беларуская", callback_data="translate_be")],
        [InlineKeyboardButton("🇩🇪 Deutsch", callback_data="translate_de"),
         InlineKeyboardButton("🇨🇳 中文", callback_data="translate_zh"),
         InlineKeyboardButton("🇸🇦 العربية", callback_data="translate_ar")],
        [InlineKeyboardButton("🇯🇵 日本語", callback_data="translate_ja")],
        [InlineKeyboardButton("◀️ Назад к результатам", callback_data="back_to_result")],
    ])

def format_time(seconds: float) -> str:
    return f"{int(seconds // 60):02d}:{int(seconds % 60):02d}"

def extract_audio(input_path: Path, output_path: Path, start_time: str = None, end_time: str = None) -> Path:
    try:
        input_kwargs = {}
        if start_time: input_kwargs["ss"] = start_time
        if end_time: input_kwargs["to"] = end_time
        (ffmpeg.input(str(input_path), **input_kwargs)
         .output(str(output_path), acodec="pcm_s16le", ac=1, ar="16000")
         .overwrite_output().run(cmd=FFMPEG_PATH, quiet=True))
        return output_path
    except ffmpeg.Error as e:
        logger.error(f"FFmpeg ошибка: {e.stderr.decode() if e.stderr else e}")
        raise

def parse_time_range(time_str: str) -> tuple:
    time_str = time_str.replace(" ", "")
    parts = re.split(r'[-–—]', time_str)
    if len(parts) != 2: raise ValueError("Неверный формат. Используй: xx:xx - xx:xx")
    for t in parts:
        if not re.match(r'^\d{1,2}:\d{2}$', t.strip()):
            raise ValueError(f"Неверный формат времени: {t}. Используй: MM:SS")
    return parts[0].strip(), parts[1].strip()

def get_audio_stats(audio_path: Path, transcript: str) -> str:
    try:
        result = subprocess.run(
            [FFPROBE_PATH, "-v", "error", "-show_entries",
             "format=duration", "-of", "default=noprint_wrappers=1:nokey=1",
             str(audio_path)], capture_output=True, text=True)
        duration_sec = float(result.stdout.strip())
        minutes, seconds = int(duration_sec // 60), int(duration_sec % 60)
        words = len(transcript.split())
        wpm = int(words / (duration_sec / 60)) if duration_sec > 0 else 0
        emoji_tempo = "🐢" if wpm < 100 else "🚶" if wpm < 160 else "🏃" if wpm < 200 else "⚡"
        label_tempo = "Медленный" if wpm < 100 else "Нормальный" if wpm < 160 else "Быстрый" if wpm < 200 else "Очень быстрый"
        return f"""
<b>📊 СТАТИСТИКА РЕЧИ</b>

⏱ <b>Длительность:</b> {minutes} мин {seconds} сек
📝 <b>Слов:</b> {words}
🗣 <b>Темп речи:</b> {wpm} слов/мин {emoji_tempo} {label_tempo}
📏 <b>Символов:</b> {len(transcript)}
"""
    except:
        return "📊 Статистика недоступна"

def transcribe_audio(audio_path: Path) -> tuple:
    global medium_model
    segments_list, segments_with_timestamps = [], []

    logger.info("🎙 Расшифровка base-моделью...")
    segments, info = whisper_model.transcribe(str(audio_path), beam_size=5)

    all_segments = list(segments)
    full_text = " ".join([s.text.strip() for s in all_segments])
    word_count = len(full_text.split())
    avg_prob = info.language_probability

    meaningful_words = [w for w in full_text.split() if len(w) > 2]
    meaningful_ratio = len(meaningful_words) / max(word_count, 1)

    quality_score = (avg_prob * 0.5) + (meaningful_ratio * 0.5)

    logger.info(f"Качество base: слов={word_count}, вероятность={avg_prob:.2f}, осмысленность={meaningful_ratio:.2f}, итог={quality_score:.2f}")

    if quality_score < 0.6 or word_count < 10:
        logger.info("⚠️ Низкое качество — загружаю medium-модель...")
        try:
            if medium_model is None:
                medium_model = WhisperModel("medium", device="cpu", compute_type="int8")
            segments_new, info = medium_model.transcribe(str(audio_path), beam_size=5)
            all_segments = list(segments_new)
            logger.info("✅ Medium-модель отработала")
        except Exception as e:
            logger.warning(f"Medium недоступна: {e}. Остаюсь на base.")

    for segment in all_segments:
        text = segment.text.strip()
        if text:
            segments_list.append(text)
            segments_with_timestamps.append({
                "start": segment.start,
                "end": segment.end,
                "text": text
            })

    final_text = " ".join(segments_list)
    logger.info(f"Итог: {len(final_text)} символов, {len(segments_list)} сегментов")

    return final_text, segments_with_timestamps, info.language

def extract_keywords_with_llm(transcript: str) -> list:
    prompt = f"""Ниже расшифровка разговора, возможно с ошибками распознавания.
Найди ОБЩИЙ СМЫСЛ, игнорируй бессмысленные слова.
Выдели 5-7 ГЛАВНЫХ ТЕМ или КЛЮЧЕВЫХ ПОНЯТИЙ (существительные или словосочетания).
Если текст неразборчив — выдели хотя бы ОДНУ вероятную тему.

ТЕКСТ: {transcript[:3000]}

ФОРМАТ: {{"keywords": ["тема1", "тема2"]}}"""

    try:
        response = llm_client.chat.completions.create(
            model="yandexgpt-lite/latest",
            messages=[{"role": "system", "content": "Ты эксперт по извлечению смысла из зашумлённых текстов. Только JSON."},
                      {"role": "user", "content": prompt}],
            temperature=0.3, max_tokens=300)
        json_match = re.search(r'\{.*\}', response.choices[0].message.content.strip(), re.DOTALL)
        if json_match:
            result = json.loads(json_match.group(0))
            keywords = result.get("keywords", [])
            if keywords: return keywords[:7]
    except Exception as e:
        logger.error(f"Ошибка ключевых слов: {e}")

    stop_words = {"это", "что", "как", "ты", "она", "там", "ещё", "уже", "только", "будешь", "много", "делать", "голова", "кстати", "почему", "когда", "говорят", "считает"}
    words = [w.strip(".,!?;:()[]{}«»\"'") for w in transcript.split()
             if len(w) > 4 and w.lower() not in stop_words]
    return [w for w, _ in Counter(words).most_common(7)]

def analyze_with_llm(transcript: str) -> dict:
    prompt = f"""Проанализируй расшифровку и выдели: темы, решения, задачи.
ТЕКСТ: {transcript[:3000]}
ФОРМАТ (строго JSON): {{"topics": ["тема"], "decisions": ["решение"], "tasks": ["задача"]}}"""
    try:
        response = llm_client.chat.completions.create(
            model="yandexgpt-lite/latest",
            messages=[{"role": "system", "content": "Отвечай только JSON."}, {"role": "user", "content": prompt}],
            temperature=0.4, max_tokens=1500)
        json_match = re.search(r'\{.*\}', response.choices[0].message.content.strip(), re.DOTALL)
        if json_match:
            result = json.loads(json_match.group(0))
            if not result.get("topics"): result["topics"] = ["Обсуждение"]
            return result
    except Exception as e:
        logger.error(f"LLM ошибка: {e}")
    return {"topics": ["Речь"], "decisions": [], "tasks": []}

def generate_timestamps(segments: list, topics: list) -> str:
    if not segments or not topics: return "⏱ Нет данных"
    total_duration = segments[-1]["end"]
    segment_duration = total_duration / len(topics)
    lines = ["<b>⏱ ТАЙМ-КОДЫ ПО ТЕМАМ</b>\n"]
    for i, topic in enumerate(topics):
        st, en = i * segment_duration, min((i + 1) * segment_duration, total_duration)
        relevant = [s for s in segments if s["start"] >= st and s["end"] <= en]
        sample = (relevant[0]["text"][:80] + "...") if relevant else "Нет текста"
        lines.append(f"<b>📌 {topic}</b>\n   ⏰ {format_time(st)} – {format_time(en)}\n   💬 «{sample}»\n")
    return "\n".join(lines)

def translate_text(text: str, target_lang: str) -> str:
    codes = {"ru": "ru", "en": "en", "be": "be", "de": "de", "zh": "zh-CN", "ar": "ar", "ja": "ja"}
    try:
        resp = requests.get("https://translate.googleapis.com/translate_a/single",
                            params={"client": "gtx", "sl": "auto", "tl": codes.get(target_lang, target_lang), "dt": "t", "q": text[:3000]}, timeout=10)
        parts = [p[0] for p in resp.json()[0] if p[0]]
        return " ".join(parts) if parts else "❌ Не удалось перевести"
    except:
        return "❌ Ошибка перевода"

def generate_mind_map(transcript: str, topics: list, decisions: list) -> BytesIO:
    keywords = extract_keywords_with_llm(transcript)
    if not keywords: keywords = topics[:7]

    G = nx.Graph()
    central = "🎯 Суть разговора"
    G.add_node(central, size=3000)

    for kw in keywords:
        G.add_node(kw, size=2000, type="keyword")
        G.add_edge(central, kw)

    for d in decisions[:3]:
        node = f"✅ {d[:30]}"
        G.add_node(node, size=1500, type="decision")
        G.add_edge(keywords[0] if keywords else central, node)

    fig, ax = plt.subplots(figsize=(14, 10), facecolor='#1e1e2e')
    ax.set_facecolor('#1e1e2e')
    pos = nx.spring_layout(G, k=3.0, iterations=100, seed=42)

    node_colors = []
    for n in G.nodes:
        if n == central: node_colors.append('#f9e2af')
        elif n.startswith("✅"): node_colors.append('#89b4fa')
        else: node_colors.append('#a6e3a1')

    node_sizes = [G.nodes[n].get("size", 1500) for n in G.nodes]

    nx.draw_networkx_nodes(G, pos, node_color=node_colors, node_size=node_sizes,
                           alpha=0.95, edgecolors='#45475a', linewidths=2, ax=ax)
    nx.draw_networkx_edges(G, pos, alpha=0.4, edge_color='#585b70', width=2, ax=ax, style='dashed')

    for node, (x, y) in pos.items():
        font_size = 13 if node == central else 10
        ax.text(x, y, node, fontsize=font_size, fontweight='bold', ha='center', va='center',
                bbox=dict(boxstyle="round,pad=0.4", facecolor='#313244', edgecolor='#45475a', alpha=0.9),
                color='#cdd6f4')

    ax.set_title("🧠 Смысловая карта разговора", fontsize=18, fontweight='bold', color='#cdd6f4', pad=20)
    ax.axis("off")
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor='#1e1e2e')
    plt.close(fig)
    buf.seek(0)
    return buf

def generate_protocol(analysis: dict, transcript: str, is_cut: bool = False,
                      start_time: str = None, end_time: str = None) -> BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_heading('ПРОТОКОЛ СОВЕЩАНИЯ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'📅 Дата: {datetime.now().strftime("%d.%m.%Y")}  🕐 Время: {datetime.now().strftime("%H:%M")}')
    if is_cut: doc.add_paragraph(f'✂️ Фрагмент: {start_time} — {end_time}')
    doc.add_paragraph('')

    doc.add_heading('1. ПОВЕСТКА ДНЯ', level=1)
    for i, t in enumerate(analysis.get('topics', []), 1): doc.add_paragraph(f'{i}. {t}', style='List Number')

    doc.add_heading('2. ПРИНЯТЫЕ РЕШЕНИЯ', level=1)
    for i, d in enumerate(analysis.get('decisions', []), 1): doc.add_paragraph(f'{i}. {d}', style='List Number')

    doc.add_heading('3. ЗАДАЧИ И ПОРУЧЕНИЯ', level=1)
    for i, t in enumerate(analysis.get('tasks', []), 1): doc.add_paragraph(f'{i}. {t}', style='List Number')

    doc.add_heading('4. ПОЛНАЯ РАСШИФРОВКА', level=1)
    doc.add_paragraph(transcript)

    doc.add_paragraph('\n' + '_' * 40)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Сгенерировано ботом Аудио-Скрайбер')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

async def process_audio_analysis(message, user_id: int, audio_path: Path, is_cut=False, ss=None, to=None):
    t0 = time.time()
    if user_id not in audio_storage: audio_storage[user_id] = str(audio_path)

    status_msg = await message.reply_text("⏳ <i>Начинаю обработку...</i>", parse_mode="HTML")

    await asyncio.sleep(0.3)
    await status_msg.edit_text("🎙 <i>Расшифровываю речь... [1/5]</i>", parse_mode="HTML")
    transcript, segments, detected_lang = transcribe_audio(audio_path)
    text_storage[user_id] = transcript
    segments_storage[user_id] = segments

    if not transcript:
        await status_msg.edit_text("❌ Не удалось распознать речь.")
        return

    lang_flags = {"ru": "🇷🇺", "en": "🇬🇧", "be": "🇧🇾", "de": "🇩🇪", "zh": "🇨🇳", "ar": "🇸🇦", "ja": "🇯🇵"}
    logger.info(f"Язык: {detected_lang}")

    await asyncio.sleep(0.3)
    await status_msg.edit_text("🧠 <i>Анализирую смысл... [2/5]</i>", parse_mode="HTML")
    analysis = analyze_with_llm(transcript)
    result_storage[user_id] = analysis
    protocol_storage[user_id] = {"analysis": analysis, "transcript": transcript, "is_cut": is_cut, "start_time": ss, "end_time": to}

    await asyncio.sleep(0.3)
    await status_msg.edit_text("📊 <i>Считаю статистику... [3/5]</i>", parse_mode="HTML")
    stats_text = get_audio_stats(audio_path, transcript)
    stats_storage[user_id] = stats_text

    await asyncio.sleep(0.3)
    await status_msg.edit_text("🖼 <i>Рисую Mind Map... [4/5]</i>", parse_mode="HTML")
    mind_map = generate_mind_map(transcript, analysis.get("topics", []), analysis.get("decisions", []))

    await asyncio.sleep(0.3)
    await status_msg.edit_text("⏱ <i>Расставляю тайм-коды... [5/5]</i>", parse_mode="HTML")
    timestamps_report = generate_timestamps(segments, analysis.get("topics", []))

    await status_msg.delete()

    header = f"✂️ <b>ОТЧЁТ (фрагмент)</b>\n{ss} – {to}" if is_cut else f"🎙 <b>ОТЧЁТ ПО ЗАПИСИ</b> {lang_flags.get(detected_lang, '🌍')}"
    topics_text = "\n".join([f"  ├ {t}" for t in analysis.get("topics", [])]) or "  └ не определены"
    decisions_text = "\n".join([f"  ├ {d}" for d in analysis.get("decisions", [])]) or "  └ не зафиксированы"
    tasks_text = "\n".join([f"  ├ {t}" for t in analysis.get("tasks", [])]) or "  └ нет"

    final_msg = f"""
{header}

<pre>╔══════════════════════════════╗</pre>
<b>🧠 ТЕМЫ ОБСУЖДЕНИЯ</b>
{topics_text}

<b>✅ ПРИНЯТЫЕ РЕШЕНИЯ</b>
{decisions_text}

<b>📋 ЗАДАЧИ</b>
{tasks_text}
<pre>╚══════════════════════════════╝</pre>

📝 <i>{transcript[:200]}...</i>
"""

    await message.reply_photo(mind_map, caption="<b>🧠 Mind Map встречи</b>", parse_mode="HTML")
    await asyncio.sleep(0.3)
    await message.reply_text(final_msg, parse_mode="HTML")
    await asyncio.sleep(0.3)
    await message.reply_text(stats_text, parse_mode="HTML")
    await asyncio.sleep(0.3)
    if len(timestamps_report) <= 4000: await message.reply_text(timestamps_report, parse_mode="HTML")
    else: await message.reply_document(BytesIO(timestamps_report.encode()), filename="timestamps.txt")

    elapsed = int(time.time() - t0)
    usage_stats["total_processed"] += 1
    usage_stats["total_seconds"] += elapsed
    await asyncio.sleep(0.3)
    await message.reply_text(f"⏱ <i>Обработано за {elapsed} сек</i>", parse_mode="HTML")

    if len(transcript) <= 4000:
        await message.reply_text(f"<b>📝 ПОЛНЫЙ ТЕКСТ</b>\n\n{transcript}", parse_mode="HTML", reply_markup=after_result_menu())
    else:
        await message.reply_document(BytesIO(transcript.encode()), filename="transcript.txt", caption="📝 Полная расшифровка")
        await message.reply_text("<b>✅ Готово!</b>", parse_mode="HTML", reply_markup=after_result_menu())

async def start_command(update: Update, context):
    text = """
<b>🎙 АУДИО-СКРАЙБЕР</b>
Твой ИИ-ассистент для работы с аудио и видео

<b>⚡ Что я умею:</b>
• Расшифровывать аудио и видео в текст
• Выделять темы, решения и задачи
• Рисовать интеллект-карту (Mind Map)
• Генерировать протокол (.docx)
• Показывать тайм-коды и статистику
• Переводить на 7 языков
• Извлекать аудиодорожку

<b>📂 Просто отправь мне файл!</b>
"""
    await update.message.reply_text(text, reply_markup=main_menu(), parse_mode="HTML")

async def button_handler(update: Update, context):
    query = update.callback_query
    await query.answer()
    uid = query.from_user.id
    d = query.data

    if d == "main_menu":
        await query.message.reply_text("✨ <b>Главное меню</b>", reply_markup=main_menu(), parse_mode="HTML")
    elif d == "about":
        await query.message.reply_text(
            "<b>ℹ️ КАК ЭТО РАБОТАЕТ</b>\n\n"
            "1️⃣ Ты отправляешь аудио/видео\n"
            "2️⃣ Whisper расшифровывает речь\n"
            "3️⃣ YandexGPT анализирует смысл\n"
            "4️⃣ Ты получаешь: отчёт, Mind Map, тайм-коды, статистику, протокол\n\n"
            "<b>🎯 Для лекций, совещаний, интервью.</b>",
            reply_markup=back_button(), parse_mode="HTML")
    elif d == "send_media":
        await query.message.reply_text(
            "📤 <b>ОТПРАВЬ ФАЙЛ</b>\n\n"
            "🎤 Голосовые сообщения\n"
            "🎵 Аудио: MP3, WAV, OGG\n"
            "🎬 Видео: MP4\n"
            "⭕ Видеокружки\n\n"
            "<i>Размер до 20 МБ</i>",
            reply_markup=back_button(), parse_mode="HTML")
    elif d == "help":
        await query.message.reply_text(
            "<b>🆘 ПОМОЩЬ</b>\n\n"
            "✂️ <b>Обрезка:</b> MM:SS - MM:SS\n"
            "📊 <b>Статистика:</b> темп, слова\n"
            "⏱ <b>Тайм-коды:</b> темы по времени\n"
            "📄 <b>Протокол:</b> готовый .docx\n"
            "🌍 <b>Перевод:</b> 7 языков\n\n"
            "/start — главное меню",
            reply_markup=back_button(), parse_mode="HTML")
    elif d == "show_stats":
        await query.message.reply_text(stats_storage.get(uid, "❌ Нет данных."), parse_mode="HTML", reply_markup=after_result_menu())
    elif d == "show_timestamps":
        if uid in segments_storage and uid in result_storage:
            ts = generate_timestamps(segments_storage[uid], result_storage[uid].get("topics", []))
            if len(ts) <= 4000: await query.message.reply_text(ts, parse_mode="HTML", reply_markup=after_result_menu())
            else: await query.message.reply_document(BytesIO(ts.encode()), filename="timestamps.txt")
        else: await query.message.reply_text("❌ Нет данных.", reply_markup=main_menu())
    elif d == "show_result":
        if uid in result_storage:
            a = result_storage[uid]
            await query.message.reply_text(
                f"<b>🎙 ОТЧЁТ</b>\n\n🧠 Темы:\n" + "\n".join(f"  ├ {t}" for t in a.get("topics", [])) +
                "\n\n✅ Решения:\n" + "\n".join(f"  ├ {d}" for d in a.get("decisions", [])) +
                "\n\n📋 Задачи:\n" + "\n".join(f"  ├ {t}" for t in a.get("tasks", [])),
                parse_mode="HTML", reply_markup=after_result_menu())
        else: await query.message.reply_text("❌ Нет данных.", reply_markup=main_menu())
    elif d == "download_protocol":
        if uid in protocol_storage:
            data = protocol_storage[uid]
            msg = await query.message.reply_text("📄 Генерирую протокол...")
            try:
                protocol = generate_protocol(data["analysis"], data["transcript"], data["is_cut"], data["start_time"], data["end_time"])
                await query.message.reply_document(protocol, filename=f"Протокол_{datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx", caption="📄 Протокол совещания")
            except Exception as e:
                await query.message.reply_text(f"❌ Ошибка: {e}")
            await msg.delete()
        else: await query.message.reply_text("❌ Нет данных.", reply_markup=main_menu())
    elif d == "download_txt":
        if uid in text_storage:
            await query.message.reply_document(BytesIO(text_storage[uid].encode()), filename="transcript.txt", caption="📝 Расшифровка")
        else: await query.message.reply_text("❌ Нет данных.", reply_markup=main_menu())
    elif d == "extract_audio_only":
        p = audio_storage.get(uid)
        if p and Path(p).exists(): await query.message.reply_audio(audio=open(p, "rb"), title="Аудиодорожка", performer="Аудио-Скрайбер")
        else: await query.message.reply_text("❌ Аудио не найдено.", reply_markup=after_result_menu())
    elif d == "cut_video":
        if uid in video_storage:
            pending_cut[uid] = True
            await query.message.reply_text("✂️ <b>Укажи время:</b> <code>MM:SS - MM:SS</code>\n\nПример: <code>01:30 - 05:45</code>", parse_mode="HTML")
        else: await query.message.reply_text("❌ Сначала отправь видео.", reply_markup=main_menu())
    elif d == "analyze_full":
        if uid in video_storage:
            inp = Path(video_storage[uid])
            msg = await query.message.reply_text("🔄 Конвертирую...")
            ap = TEMP_DIR / f"full_{uid}_{int(time.time())}.wav"
            extract_audio(inp, ap)
            audio_storage[uid] = str(ap)
            await msg.delete()
            await process_audio_analysis(query.message, uid, ap)
        else: await query.message.reply_text("❌ Видео не найдено.", reply_markup=main_menu())
    elif d == "translate_menu":
        await query.message.reply_text("🌍 <b>Выбери язык:</b>", reply_markup=translate_menu(), parse_mode="HTML")
    elif d == "back_to_result":
        await query.message.reply_text("📋 <b>Действия:</b>", reply_markup=after_result_menu(), parse_mode="HTML")
    elif d.startswith("translate_"):
        lang = d.replace("translate_", "")
        names = {"ru": "Русский", "en": "English", "be": "Беларуская", "de": "Deutsch", "zh": "中文", "ar": "العربية", "ja": "日本語"}
        if uid in text_storage:
            tr = translate_text(text_storage[uid], lang)
            await query.message.reply_text(f"🌍 <b>{names.get(lang, lang)}</b>\n\n{tr}", parse_mode="HTML", reply_markup=translate_menu())
        else: await query.message.reply_text("❌ Текст не найден.", reply_markup=main_menu())

async def handle_text(update: Update, context):
    uid = update.message.from_user.id
    if pending_cut.get(uid) and uid in video_storage:
        try:
            ss, to = parse_time_range(update.message.text.strip())
            pending_cut[uid] = False
            ap = TEMP_DIR / f"cut_{update.message.message_id}.wav"
            extract_audio(Path(video_storage[uid]), ap, ss, to)
            await process_audio_analysis(update.message, uid, ap, True, ss, to)
        except ValueError as e:
            await update.message.reply_text(f"❌ {e}\nФормат: <code>MM:SS - MM:SS</code>", parse_mode="HTML")

async def handle_media(update: Update, context):
    m = update.message
    uid = m.from_user.id

    if m.voice: fid, suf, mt = m.voice.file_id, "voice.ogg", "🎤 Голосовое сообщение"
    elif m.audio: fid, suf, mt = m.audio.file_id, "audio.mp3", "🎵 Аудиофайл"
    elif m.video: fid, suf, mt = m.video.file_id, "video.mp4", "🎬 Видеофайл"
    elif m.video_note: fid, suf, mt = m.video_note.file_id, "video_note.mp4", "⭕ Видеокружок"
    else: return

    file_obj = m.voice or m.audio or m.video or m.video_note
    if hasattr(file_obj, 'file_size'):
        size_mb = file_obj.file_size / (1024 * 1024)
        if size_mb > 20:
            await m.reply_text(f"❌ <b>Файл слишком большой:</b> {size_mb:.1f} МБ\n\n<i>Максимальный размер: 20 МБ</i>", parse_mode="HTML")
            return

    await m.reply_text(f"📥 <b>{mt}</b> получен!\n<i>Выбери действие:</i>", reply_markup=cut_or_full_menu(), parse_mode="HTML")
    inp = TEMP_DIR / f"{m.message_id}_{suf}"
    try:
        await (await context.bot.get_file(fid)).download_to_drive(inp)
        video_storage[uid] = str(inp)
        if mt in ["🎤 Голосовое сообщение", "🎵 Аудиофайл"]: audio_storage[uid] = str(inp)
    except Exception as e:
        await m.reply_text(f"❌ Ошибка при скачивании: {e}")

def main():
    logger.info("Запуск бота...")
    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", start_command))
    app.add_handler(CallbackQueryHandler(button_handler))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(MessageHandler(filters.VOICE | filters.AUDIO | filters.VIDEO | filters.VIDEO_NOTE, handle_media))
    logger.info("Бот запущен!")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()